"""
parser.py
---------
Handles raw text extraction from PDF and DOCX files.
"""

import pdfplumber
import docx
import io


def extract_text_from_pdf(file) -> str:
    """Extract all text from an uploaded PDF file object."""
    text = ""
    try:
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {e}")
    return text.strip()


def extract_text_from_docx(file) -> str:
    """Extract all text from an uploaded DOCX file object."""
    text = ""
    try:
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {e}")
    return text.strip()


def extract_text(file, filename: str) -> str:
    """
    Detect file type by extension and call appropriate extractor.
    Returns plain text string.
    """
    ext = filename.lower().split(".")[-1]
    if ext == "pdf":
        return extract_text_from_pdf(file)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file)
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Please upload PDF or DOCX.")
